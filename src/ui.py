"""
Módulo de Interfaz de Usuario para los 5 Pasos Operativos, Carga Masiva,
Gestión de Usuarios, Dashboard, Consolidado General, Interfaz Segura y Alertas Sanitarias.
"""
import io
import os
import time
import sqlite3
import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.database import get_connection
from src.catalogo import (
    get_catalogo, get_opciones_selectbox, buscar_por_etiqueta,
    BODEGAS_OFICIALES, PROVEEDORES_OFICIALES, TIPOS_DOCUMENTO, OPCION_MANUAL
)
from src.importer import procesar_carga_masiva

ESTADOS_TRAMITE_PROVEEDOR = [
    "Correo enviado - En espera de respuesta de proveedor",
    "Correo enviado - Canje aceptado",
    "Correo enviado - Canje rechazado",
    "Correo enviado - En espera de retiro por proveedor",
    "Bulto retirado por proveedor",
    "Canje aceptado - En espera de Nota de credito",
    "Canje aceptado - En espera de reposición del producto",
    "Nota de credito recibida",
    "Producto recibido"
]

BODEGAS_PASO4 = BODEGAS_OFICIALES + ["Bodega de Excluidos"]
OPCIONES_FISICA_P4 = BODEGAS_PASO4 + ["Despachado / Retirado de Bodega"]

OPCIONES_DIFUSION_RED = [
    "No iniciada la difusión",
    "Ofrecido a la Red",
    "Aceptado por Establecimiento",
    "Producto donado"
]

OPCIONES_REDISTRIBUCION_STOCK = [
    "Baja mensual CENABAST",
    "Reprogramación anual CENABAST",
    "Pedido especial CENABAST",
    "Consulta de compra"
]

TEMAS_COLOR_CLAROS = {
    "Claro (Predeterminado)": {"bg": "#f8f9fa", "card": "#ffffff", "text": "#1a1d20", "border": "#dee2e6"},
    "Azul Pastel": {"bg": "#edf2f7", "card": "#e2e8f0", "text": "#0f172a", "border": "#cbd5e1"},
    "Rojo Suave": {"bg": "#fef2f2", "card": "#fee2e2", "text": "#450a0a", "border": "#fca5a5"},
    "Amarillo / Crema": {"bg": "#fefce8", "card": "#fef08a", "text": "#422006", "border": "#fde047"},
    "Violeta Lavanda": {"bg": "#faf5ff", "card": "#f3e8ff", "text": "#3b0764", "border": "#d8b4fe"},
    "Verde Menta": {"bg": "#f0fdf4", "card": "#dcfce7", "text": "#052e16", "border": "#86efac"},
    "Rosado Pastel": {"bg": "#fdf2f8", "card": "#fce7f3", "text": "#500724", "border": "#fbcfe8"}
}

def aplicar_estilo_tema(nombre_tema):
    tema = TEMAS_COLOR_CLAROS.get(nombre_tema, TEMAS_COLOR_CLAROS["Claro (Predeterminado)"])
    css = f"""
    <style>
        [data-testid="stHeader"] {{ display: none !important; }}
        #MainMenu {{ visibility: hidden !important; }}
        footer {{ visibility: hidden !important; }}
        .stDeployButton {{ display: none !important; }}
        .block-container {{ padding-top: 1.5rem !important; padding-bottom: 1.5rem !important; }}
        .stApp {{ background-color: {tema['bg']} !important; color: {tema['text']} !important; }}
        [data-testid="stSidebar"] {{ background-color: {tema['card']} !important; border-right: 1px solid {tema['border']} !important; }}
        [data-testid="stSidebar"] *, [data-testid="stSidebar"] label, [data-testid="stSidebar"] p, [data-testid="stSidebar"] span {{ color: {tema['text']} !important; font-weight: 600 !important; }}
        [data-testid="stSidebar"] div[role="radiogroup"] > label {{ padding: 10px 12px; background-color: transparent; border-radius: 8px; margin-bottom: 4px; transition: all 0.2s ease; }}
        [data-testid="stSidebar"] div[role="radiogroup"] > label:hover {{ background-color: rgba(0, 0, 0, 0.04); transform: translateX(4px); }}
        [data-testid="stSidebar"] div[role="radiogroup"] > label > div:first-child {{ transform: scale(0.85); }}
        .stMarkdown, .stText, h1, h2, h3, h4, h5, h6, label, p, span {{ color: {tema['text']} !important; }}
        [data-testid="stMetricValue"] {{ color: {tema['text']} !important; }}
        .stMarkdown p {{ margin-bottom: 0 !important; }}
        code {{ background-color: #dcfce7 !important; color: #15803d !important; border: 1px solid #86efac !important; font-weight: bold !important; padding: 2px 8px !important; border-radius: 4px !important; }}
        [data-testid="collapsedControl"], [data-testid="stSidebarCollapsedControl"] {{ opacity: 1 !important; background-color: #ffffff !important; border: 1px solid #cbd5e1 !important; border-radius: 8px !important; box-shadow: 0 2px 5px rgba(0,0,0,0.08) !important; color: #0f172a !important; transition: all 0.3s ease; z-index: 999999 !important; }}
        [data-testid="collapsedControl"]:hover, [data-testid="stSidebarCollapsedControl"]:hover {{ background-color: #f1f5f9 !important; border-color: #94a3b8 !important; }}
        [data-testid="collapsedControl"] svg, [data-testid="stSidebarCollapsedControl"] svg {{ fill: #0f172a !important; color: #0f172a !important; }}
        div[data-baseweb="select"] > div, .stTextInput div[data-baseweb="input"], .stNumberInput div[data-baseweb="input"], .stDateInput div[data-baseweb="input"], .stTextArea div[data-baseweb="textarea"], div[data-baseweb="popover"] {{ background-color: #ffffff !important; border: 1px solid #cbd5e1 !important; }}
        div[data-baseweb="select"] * {{ color: #0f172a !important; }}
        .stTextInput input, .stNumberInput input, .stDateInput input, .stTextArea textarea {{ background-color: #ffffff !important; color: #0f172a !important; }}
        ul[data-testid="stSelectboxVirtualDropdown"] {{ background-color: #ffffff !important; border: 1px solid #cbd5e1 !important; }}
        ul[data-testid="stSelectboxVirtualDropdown"] li {{ background-color: #ffffff !important; color: #0f172a !important; }}
        ul[data-testid="stSelectboxVirtualDropdown"] li:hover {{ background-color: #e2e8f0 !important; color: #0f172a !important; }}
        .stButton button, .stDownloadButton button, [data-testid="stFileUploader"] button {{ background-color: #ffffff !important; color: #0f172a !important; border: 1px solid #cbd5e1 !important; font-weight: bold !important; box-shadow: 0 1px 2px rgba(0,0,0,0.05) !important; }}
        .stButton button *, .stDownloadButton button *, [data-testid="stFileUploader"] button * {{ color: #0f172a !important; }}
        div[data-testid="stForm"] button {{ background-color: #0284c7 !important; border: 1px solid #0369a1 !important; }}
        div[data-testid="stForm"] button *, div[data-testid="stForm"] button p, div[data-testid="stForm"] button span {{ color: #ffffff !important; font-weight: bold !important; font-size: 15px !important; }}
        div[data-testid="stForm"] button:hover {{ background-color: #0369a1 !important; }}
        [data-testid="stFileUploader"] section {{ background-color: #ffffff !important; border: 2px dashed #cbd5e1 !important; }}
        [data-testid="stFileUploader"] section * {{ color: #334155 !important; }}
        [data-testid="stSidebar"] [data-testid="stImage"] {{ display: flex !important; justify-content: center !important; margin-bottom: 20px !important; background: transparent !important; border: none !important; box-shadow: none !important; padding: 0 !important; }}
        [data-testid="stSidebar"] [data-testid="stImage"] img {{ display: block !important; margin: 0 auto !important; image-rendering: high-quality !important; -webkit-font-smoothing: antialiased !important; mix-blend-mode: multiply !important; }}
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

def calcular_semaforo_vencimiento(fecha_str, motivo):
    """Calcula el nivel de alerta según los días restantes."""
    if motivo == "Alerta Sanitaria":
        return "🚨 ALERTA (ISP)"
    if not fecha_str or pd.isna(fecha_str): 
        return "⚪ N/A"
    try:
        venc = datetime.strptime(str(fecha_str), "%Y-%m-%d").date()
        dias = (venc - datetime.now().date()).days
        if dias <= 60:
            return "🔴 Roja (≤ 60d)"
        elif dias <= 120:
            return "🟡 Amarilla (61-120d)"
        else:
            return "🟢 Verde (> 120d)"
    except:
        return "⚪ Error Fecha"

def generar_anexo_ii_docx(datos):
    doc = Document()
    p_membrete = doc.add_paragraph()
    run_mem1 = p_membrete.add_run("DEPARTAMENTO AGENCIA NACIONAL DE MEDICAMENTOS\n")
    run_mem2 = p_membrete.add_run("SUBDEPARTAMENTO DE FISCALIZACIÓN")
    run_mem1.font.size = Pt(8); run_mem1.font.bold = True; run_mem1.font.color.rgb = RGBColor(128, 128, 128)
    run_mem2.font.size = Pt(8); run_mem2.font.bold = True; run_mem2.font.color.rgb = RGBColor(128, 128, 128)
    p_titulo = doc.add_paragraph(); p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t1 = p_titulo.add_run("ANEXO II\n"); run_t2 = p_titulo.add_run("FORMULARIO DE REPORTE DE EXISTENCIAS DE PRODUCTO RETIRADO DEL MERCADO\n"); run_t3 = p_titulo.add_run("PRODUCTOS FARMACÉUTICOS (ARTS. 60° y 71° 3, D.S. N° 3/2010)\n")
    for r in [run_t1, run_t2, run_t3]: r.font.bold = True; r.font.size = Pt(11); r.font.name = 'Arial'
    p_intro = doc.add_paragraph(); p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_intro = p_intro.add_run("Documento emitido por cada establecimiento que ha recibido producto afecto(s) a retiro del mercado, que da cuenta de la cantidad recibida, stock y cantidad distribuida de los lotes en cuestión, que debe ser informado al distribuidor de quien obtuvo el producto, mediante el presente formato o por otro documento que contenga la misma información, para ser reportado finalmente al titular del producto sujeto a retiro del mercado.")
    run_intro.font.italic = True; run_intro.font.size = Pt(11); doc.add_paragraph() 
    p_h1 = doc.add_paragraph(); r_h1 = p_h1.add_run("ANTECEDENTES DEL PRODUCTO EN PROCESO DE RETIRO DEL MERCADO\n"); r_h1.font.bold = True
    r_h1_sub = p_h1.add_run("(Datos aportados por el titular o distribuidor al establecimiento receptor que suscribe la información del presente formulario)")
    r_h1_sub.font.bold = True; r_h1_sub.font.italic = True
    t1 = doc.add_table(rows=5, cols=2); t1.style = 'Table Grid'
    t1.cell(0,0).text = "PRODUCTO"; t1.cell(0,1).text = datos.get('descripcion', '')
    t1.cell(1,0).text = "PRINCIPIO ACTIVO"; t1.cell(1,1).text = datos.get('principio_activo', '')
    t1.cell(2,0).text = "TITULAR"; t1.cell(2,1).text = datos.get('titular', '')
    t1.cell(3,0).text = "N° DE REGISTRO SANITARIO"; t1.cell(3,1).text = datos.get('registro_sanitario', '')
    t1.cell(4,0).text = "NUMERO DE SERIE(S) / LOTE(S)"; t1.cell(4,1).text = datos.get('lote', '')
    doc.add_paragraph() 
    p_h2 = doc.add_paragraph(); r_h2 = p_h2.add_run("ANTECEDENTES DEL ESTABLECIMIENTO RECEPTOR QUE SUSCRIBE LA INFORMACIÓN"); r_h2.font.bold = True
    t2 = doc.add_table(rows=5, cols=2); t2.style = 'Table Grid'
    t2.cell(0,0).text = "NOMBRE y DIRECCIÓN"; t2.cell(0,1).text = "Hospital Penco Lirquén"
    t2.cell(1,0).text = "TIPO DE ESTABLECIMIENTO/LÍNEAS DE ACTIVIDAD"; t2.cell(1,1).text = "Establecimiento de Salud Pública"
    t2.cell(2,0).text = "REPRESENTANTE LEGAL"; t2.cell(2,1).text = datos.get('representante_legal', '')
    t2.cell(3,0).text = "DIRECTOR TÉCNICO/RESPONSABLE TÉCNICO"; t2.cell(3,1).text = datos.get('director_tecnico', '')
    t2.cell(4,0).text = "FECHA DE REPORTE"; t2.cell(4,1).text = datetime.now().strftime('%d-%m-%Y')
    doc.add_paragraph() 
    p_h3 = doc.add_paragraph(); r_h3 = p_h3.add_run("ANTECEDENTES MOVIMIENTO DEL PRODUCTO"); r_h3.font.bold = True
    t3 = doc.add_table(rows=8, cols=2); t3.style = 'Table Grid'
    t3.cell(0,0).text = "PROVEEDOR"; t3.cell(0,1).text = datos.get('proveedor', '')
    t3.cell(1,0).text = "CANTIDAD RECIBIDA EN EL ESTABLECIMIENTO RECEPTOR"; t3.cell(1,1).text = "N/A"
    t3.cell(2,0).text = "CANTIDAD EN EXISTENCIA (STOCK) NO DISTRIBUIDA"; t3.cell(2,1).text = f"{datos.get('cantidad', '')} {datos.get('unidad', '')}"
    t3.cell(3,0).text = "CANTIDAD DISTRIBUIDA A OTROS ESTABLECIMIENTOS RECEPTORES"; t3.cell(3,1).text = "0"
    t3.cell(4,0).text = "REGISTRO DE DISTRIBUCIÓN"; t3.cell(4,1).text = "N/A"
    t3.cell(5,0).text = "OTRAS OBSERVACIONES"; t3.cell(5,1).text = datos.get('observaciones', '')
    t3.cell(6,0).text = "IDENTIFICACIÓN DE NOTIFICANTE y TELÉFONO DIRECTO"; t3.cell(6,1).text = f"{datos.get('director_tecnico', '')} - "
    t3.cell(7,0).text = "FIRMA DE NOTIFICANTE"; t3.cell(7,1).text = "" 
    bio = io.BytesIO(); doc.save(bio)
    return bio.getvalue()

def actualizar_bd_alertas(conn):
    cursor = conn.cursor()
    columnas = ["alerta_numero TEXT", "alerta_fecha TEXT", "titular_registro TEXT", "registro_sanitario TEXT", "principio_activo TEXT"]
    for col in columnas:
        try: cursor.execute(f"ALTER TABLE productos ADD COLUMN {col}")
        except sqlite3.OperationalError: pass 
    conn.commit()

def render_ui(user_info: dict):
    conn = get_connection()
    actualizar_bd_alertas(conn)
    
    opciones_temas = list(TEMAS_COLOR_CLAROS.keys())
    tema_actual = st.session_state.get("tema_seleccionado", "Claro (Predeterminado)")
    if tema_actual not in opciones_temas: tema_actual = "Claro (Predeterminado)"
    idx_tema = opciones_temas.index(tema_actual)
    aplicar_estilo_tema(tema_actual)

    col_user, col_rol, col_btn = st.columns([5, 4, 2], vertical_alignment="center")
    with col_user: st.markdown(f"👤 **Usuario:** {user_info['nombre_completo']}")
    with col_rol: st.markdown(f"🛡️ **Rol:** `{user_info['rol'].upper()}`")
    with col_btn:
        if st.button("🚪 Cerrar Sesión", use_container_width=True): st.session_state["logged_in"] = False; st.rerun()

    borde_color = TEMAS_COLOR_CLAROS[tema_actual]["border"]
    st.markdown(f"<hr style='margin-top: 0.5rem; margin-bottom: 1rem; border: none; border-top: 1px solid {borde_color};' />", unsafe_allow_html=True)

    path1, path2 = "assets/hospital-penco-lirquen.png", "assets/logo.png"
    if os.path.exists(path1): st.sidebar.image(path1, width=150)
    elif os.path.exists(path2): st.sidebar.image(path2, width=150)

    rol = user_info['rol']
    
    st.sidebar.markdown("### 📂 SELECCIONE MÓDULO")
    opciones_modulos = ["💊 Gestión de Vencimientos", "🚨 Alertas Sanitarias", "⚙️ Reportes y Adm."]
    modulo_sel = st.sidebar.selectbox("Módulos Principales", opciones_modulos, label_visibility="collapsed")
    st.sidebar.markdown(f"<hr style='margin: 0.5rem 0; border: none; border-top: 1px solid {borde_color};' />", unsafe_allow_html=True)
    st.sidebar.markdown(f"<p style='font-size: 14px; margin-bottom: 8px;'><b>Estás en:</b> {modulo_sel.split(' ', 1)[1]}</p>", unsafe_allow_html=True)
    
    tabs_disponibles = []
    if modulo_sel == "💊 Gestión de Vencimientos":
        if rol in ["admin", "bodega"]:
            tabs_disponibles.extend(["📋 1. Informe Bodega", "📤 Carga Masiva"])
        if rol in ["admin", "jefatura"]: tabs_disponibles.append("⚖️ 2. Canjes (Jefatura)")
        if rol in ["admin", "registro"]: tabs_disponibles.append("🚚 3. Registro/Prov.")
        if rol in ["admin", "bodega"]: tabs_disponibles.append("📦 4. Bulto/Ubicación")
        if rol in ["admin", "jefatura"]: tabs_disponibles.append("📜 5. Resolución/Cierre")
    elif modulo_sel == "🚨 Alertas Sanitarias":
        if rol in ["admin", "bodega"]: tabs_disponibles.append("📋 1. Ingresar Nueva Alerta")
        if rol in ["admin", "jefatura"]: tabs_disponibles.append("🚨 Gestión Anexo II (Jefatura)")
    elif modulo_sel == "⚙️ Reportes y Adm.":
        tabs_disponibles.extend(["🔍 Consolidado General", "📊 Dashboard / Análisis"])
        if rol == "admin": tabs_disponibles.append("👥 Gestión de Usuarios")

    tab_seleccionada = st.sidebar.radio("Pasos", tabs_disponibles, label_visibility="collapsed")

    st.sidebar.markdown(f"<hr style='margin: 1.5rem 0 0.5rem 0; border: none; border-top: 1px dashed {borde_color};' />", unsafe_allow_html=True)
    st.sidebar.markdown("<p style='font-size: 13px; margin-bottom: 4px;'>🎨 Tema Visual</p>", unsafe_allow_html=True)
    tema_sel = st.sidebar.selectbox("Tema Visual", opciones_temas, index=idx_tema, label_visibility="collapsed", key="tema_sidebar")
    if tema_sel != tema_actual: st.session_state["tema_seleccionado"] = tema_sel; st.rerun()

    # --- PASO 1 ---
    if tab_seleccionada in ["📋 1. Informe Bodega", "📋 1. Ingresar Nueva Alerta"]:
        es_modulo_alerta = (tab_seleccionada == "📋 1. Ingresar Nueva Alerta")
        if es_modulo_alerta: st.header("🚨 Ingreso Rápido de Alerta Sanitaria"); st.caption("Los productos pasarán a Cuarentena inmediatamente.")
        else: st.header("📋 Paso 1 — Informe de Bodega")
            
        catalogo = get_catalogo()
        opciones = get_opciones_selectbox(catalogo)
        
        tab_p1_ingresar, tab_p1_editar = st.tabs(["➕ Nuevo Registro", "✏️ Editar / Modificar Registros Existentes"])

        with tab_p1_ingresar:
            c_busq1, c_busq2 = st.columns(2)
            with c_busq1:
                bodega = st.selectbox("Bodega/Farmacia Origen *", BODEGAS_OFICIALES, key="b_ingreso")
                tipo_prod = st.selectbox("Tipo de producto", ["Fármaco", "Insumo"], key="tp_ingreso")
            with c_busq2:
                prod_sel = st.selectbox("Buscar Código Reyimen / Producto *", opciones, key="busqueda_producto")
                
            cod_auto, desc_auto, unidad_auto = "", "", ""
            item = buscar_por_etiqueta(catalogo, prod_sel)
            if item: cod_auto = item.get("codigo", ""); desc_auto = item.get("descripcion", ""); unidad_auto = item.get("unidad", "")
            st.divider()

            if es_modulo_alerta:
                motivo = "Alerta Sanitaria"
                st.error("🚨 MODO ALERTA SANITARIA: Se han activado los campos de urgencia.")
            else:
                motivo = "Gestión pronto vencimiento"

            with st.form("form_paso1"):
                c1, c2 = st.columns(2)
                codigo = c1.text_input("Código Reyimen *", value=cod_auto)
                unidad = c2.text_input("Unidad *", value=unidad_auto)
                descripcion = c1.text_input("Descripción *", value=desc_auto)
                cantidad = c2.number_input("Cantidad *", min_value=0.0, step=1.0)
                
                if es_modulo_alerta:
                    tipo_compra = "No Aplica (Alerta)"
                    alerta_numero = c1.text_input("N° Alerta Sanitaria ISP *")
                    num_bulto_alerta = c2.text_input("Número de Bulto Físico (Bodega Excluidos) *")
                    alerta_fecha = c1.date_input("Fecha de Alerta ISP *")
                    vencimiento = c2.date_input("Fecha de Vencimiento *")
                    lote = c1.text_input("Lote *")
                else:
                    tipo_compra = c1.selectbox("Tipo de compra *", ["CENABAST", "Compra propia"])
                    vencimiento = c2.date_input("Fecha de Vencimiento *")
                    lote = c1.text_input("Lote *")
                
                if st.form_submit_button("Guardar Paso 1"):
                    if not codigo or not descripcion or not lote or (es_modulo_alerta and (not alerta_numero or not num_bulto_alerta)):
                        st.error("Complete todos los campos obligatorios (*)")
                    else:
                        conn.cursor().execute("SELECT id FROM productos WHERE codigo_reyimen=? AND lote=? AND bodega_origen=? AND estado_global IN ('En trámite', 'CUARENTENA')", (codigo, lote, bodega))
                        if conn.cursor().fetchone(): st.warning("⚠️ Este producto ya fue ingresado y está activo.")
                        else:
                            estado_inicial = 'CUARENTENA' if es_modulo_alerta else 'En trámite'
                            ub_fisica = "Bodega de Excluidos" if es_modulo_alerta else ""
                            bulto = num_bulto_alerta if es_modulo_alerta else ""
                            a_num = alerta_numero if es_modulo_alerta else ""
                            a_fec = str(alerta_fecha) if es_modulo_alerta else ""
                            conn.cursor().execute("""
                            INSERT INTO productos (bodega_origen, tipo_producto, codigo_reyimen, descripcion, unidad, cantidad, vencimiento, lote, motivo_informe, tipo_documento, usuario_registro, paso_actual, estado_global, ubicacion_fisica, numero_bulto, alerta_numero, alerta_fecha)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                            """, (bodega, tipo_prod, codigo, descripcion, unidad, cantidad, str(vencimiento), lote, motivo, tipo_compra, user_info['usuario'], 2, estado_inicial, ub_fisica, bulto, a_num, a_fec))
                            conn.commit()
                            st.success("✅ Producto registrado."); time.sleep(1.5); st.rerun()

        with tab_p1_editar:
            st.subheader("Registros Ingresados en Paso 1 (Modificables)")
            df_p1 = pd.read_sql_query("SELECT id AS ID, bodega_origen AS Bodega, codigo_reyimen AS Código, descripcion AS Descripción, lote AS Lote, motivo_informe AS Motivo, estado_global AS Estado FROM productos WHERE paso_actual IN (1, 2) OR estado_global = 'CUARENTENA'", conn)
            if df_p1.empty: st.info("No hay registros.")
            else:
                st.dataframe(df_p1, hide_index=True, use_container_width=True, height=180)
                id_mod = st.selectbox("Seleccione ID", df_p1['ID'].tolist())
                prod_data = conn.cursor().execute("SELECT * FROM productos WHERE id=?", (id_mod,)).fetchone()
                if prod_data:
                    with st.form("form_editar_p1"):
                        st.markdown(f"**Editando ID #{id_mod}**")
                        c_e1, c_e2 = st.columns(2)
                        idx_bodega = BODEGAS_OFICIALES.index(prod_data['bodega_origen']) if prod_data['bodega_origen'] in BODEGAS_OFICIALES else 0
                        new_bodega = c_e1.selectbox("Bodega Origen *", BODEGAS_OFICIALES, index=idx_bodega)
                        new_cantidad = c_e2.number_input("Cantidad *", value=float(prod_data['cantidad']), min_value=0.0, step=1.0)
                        new_lote = c_e1.text_input("Lote *", value=prod_data['lote'])
                        try: fecha_init = datetime.strptime(prod_data['vencimiento'], "%Y-%m-%d").date()
                        except: fecha_init = datetime.now().date()
                        new_vencimiento = c_e2.date_input("Fecha de Vencimiento *", value=fecha_init)
                        b1, b2 = st.columns(2)
                        if b1.form_submit_button("💾 Guardar"):
                            conn.cursor().execute("UPDATE productos SET bodega_origen=?, cantidad=?, lote=?, vencimiento=? WHERE id=?", (new_bodega, new_cantidad, new_lote, str(new_vencimiento), id_mod))
                            conn.commit(); st.success("Actualizado."); time.sleep(1); st.rerun()
                        if b2.form_submit_button("🗑️ Eliminar"):
                            conn.cursor().execute("DELETE FROM productos WHERE id=?", (id_mod,))
                            conn.commit(); st.warning("Eliminado."); time.sleep(1); st.rerun()

    # --- PASO 2 ---
    elif tab_seleccionada == "⚖️ 2. Canjes (Jefatura)":
        st.header("⚖️ Paso 2 — Gestión de Canjes (Jefatura)")
        df = pd.read_sql_query("SELECT id AS ID, bodega_origen AS Bodega, codigo_reyimen AS Código, descripcion AS Descripción, tipo_documento AS Compra, cantidad AS Cant, lote AS Lote, vencimiento AS Vencimiento, motivo_informe AS Motivo FROM productos WHERE paso_actual = 2 AND estado_global = 'En trámite' AND motivo_informe != 'Alerta Sanitaria'", conn)
        
        if df.empty: 
            st.info("No hay productos pendientes de canje comercial.")
        else:
            hoy = datetime.now().date()
            df["Meses Vencer"] = df["Vencimiento"].apply(lambda v: max(0.0, round((datetime.strptime(str(v), "%Y-%m-%d").date() - hoy).days / 30.44, 1)) if pd.notnull(v) else 0.0)
            df["Nivel Alerta"] = df.apply(lambda row: calcular_semaforo_vencimiento(row["Vencimiento"], row["Motivo"]), axis=1)
            columnas_orden = ["ID", "Nivel Alerta", "Código", "Descripción", "Bodega", "Compra", "Cant", "Lote", "Vencimiento"]
            df = df[columnas_orden]
            
            st.markdown("### 🔍 Filtros de Búsqueda")
            col_f1, col_f2 = st.columns(2)
            filtro_semaforo = col_f1.multiselect("Filtrar por Semáforo de Riesgo", ["🔴 Roja (≤ 60d)", "🟡 Amarilla (61-120d)", "🟢 Verde (> 120d)"], placeholder="Seleccione colores...", key="sem_p2")
            filtro_codigo = col_f2.text_input("Filtrar por Código Reyimen", placeholder="Ej: 543...", key="cod_p2")
            
            df_filtrado = df.copy()
            if filtro_semaforo: df_filtrado = df_filtrado[df_filtrado["Nivel Alerta"].isin(filtro_semaforo)]
            if filtro_codigo: df_filtrado = df_filtrado[df_filtrado["Código"].astype(str).str.contains(filtro_codigo, case=False, na=False)]
            
            st.markdown("<br>", unsafe_allow_html=True)
            if df_filtrado.empty:
                st.warning("No hay productos pendientes que coincidan con los filtros seleccionados.")
            else:
                cols_mostrar = ["ID", "Nivel Alerta", "Código", "Descripción", "Bodega", "Compra", "Cant", "Lote"]
                st.dataframe(df_filtrado[cols_mostrar], hide_index=True, use_container_width=True, height=180)
                
                opciones_id = df_filtrado['ID'].tolist()
                formato_opciones = {row['ID']: f"{row['Código']} - {row['Descripción']} (Lote: {row['Lote']})" for idx, row in df_filtrado.iterrows()}
                
                st.markdown("### ⚙️ Gestión del Producto")
                prod_id = st.selectbox("Seleccione el producto filtrado que desea gestionar:", opciones_id, format_func=lambda x: formato_opciones[x], key="sel_p2")
                
                with st.form("form_paso2"):
                    aplica_canje = st.selectbox("¿Aplica Canje? *", ["Aplica", "No aplica"])
                    obs_jefatura = st.text_area("Observaciones (Normas del proveedor, exigencias, etc.)")
                    if st.form_submit_button("Avanzar a Paso 3"):
                        conn.cursor().execute("UPDATE productos SET estado_canje=?, observacion_paso2=?, fecha_paso2=?, paso_actual=3 WHERE id=?", (aplica_canje, obs_jefatura, str(datetime.now().date()), prod_id))
                        conn.commit(); st.success("Avanzado."); time.sleep(1.5); st.rerun()

    # --- PASO 3 ---
    elif tab_seleccionada == "🚚 3. Registro/Prov.":
        st.header("🚚 Paso 3 — Área de Registro / Proveedor")
        tab_p3_nuevo, tab_p3_seguimiento = st.tabs(["➕ Pendientes de Ingreso", "🔄 Seguimiento de Trámites"])
        
        with tab_p3_nuevo:
            df_p3_nuevo = pd.read_sql_query("SELECT id AS ID, codigo_reyimen AS Código, descripcion AS Descripción, cantidad AS Cant, lote AS Lote, vencimiento AS Vencimiento, motivo_informe AS Motivo, estado_canje AS Canje, observacion_paso2 AS Obs_P2 FROM productos WHERE paso_actual = 3 AND estado_global = 'En trámite'", conn)
            
            if df_p3_nuevo.empty: 
                st.info("No hay nuevos productos.")
            else:
                df_p3_nuevo["Nivel Alerta"] = df_p3_nuevo.apply(lambda row: calcular_semaforo_vencimiento(row["Vencimiento"], row["Motivo"]), axis=1)
                
                st.markdown("### 🔍 Filtros de Búsqueda")
                col_f1, col_f2 = st.columns(2)
                filtro_semaforo = col_f1.multiselect("Filtrar por Semáforo de Riesgo", ["🔴 Roja (≤ 60d)", "🟡 Amarilla (61-120d)", "🟢 Verde (> 120d)"], placeholder="Seleccione colores...", key="sem_p3_n")
                filtro_codigo = col_f2.text_input("Filtrar por Código Reyimen", placeholder="Ej: 543...", key="cod_p3_n")
                
                df_filtrado = df_p3_nuevo.copy()
                if filtro_semaforo: df_filtrado = df_filtrado[df_filtrado["Nivel Alerta"].isin(filtro_semaforo)]
                if filtro_codigo: df_filtrado = df_filtrado[df_filtrado["Código"].astype(str).str.contains(filtro_codigo, case=False, na=False)]
                
                st.markdown("<br>", unsafe_allow_html=True)
                if df_filtrado.empty:
                    st.warning("No hay productos pendientes que coincidan con los filtros.")
                else:
                    cols_mostrar = ["ID", "Nivel Alerta", "Código", "Descripción", "Cant", "Lote", "Canje"]
                    st.dataframe(df_filtrado[cols_mostrar], hide_index=True, use_container_width=True, height=180)
                    
                    opciones_id = df_filtrado['ID'].tolist()
                    formato_opciones = {row['ID']: f"{row['Código']} - {row['Descripción']} (Lote: {row['Lote']})" for idx, row in df_filtrado.iterrows()}
                    
                    st.markdown("### ⚙️ Gestión del Trámite")
                    prod_id = st.selectbox("Seleccione el producto filtrado para ingresar trámite:", opciones_id, format_func=lambda x: formato_opciones[x], key="sel_p3_n")
                    
                    obs_previa = df_filtrado.loc[df_filtrado['ID'] == prod_id, 'Obs_P2'].values[0]
                    if obs_previa: st.info(f"📝 **Instrucciones de Jefatura:** {obs_previa}")
                    
                    with st.form("form_paso3_ingreso"):
                        proveedor, tipo_doc = st.selectbox("Proveedor *", PROVEEDORES_OFICIALES), st.selectbox("Tipo Doc *", TIPOS_DOCUMENTO)
                        num_doc, tramite = st.text_input("N° Documento / OC *"), st.selectbox("Estado del trámite *", ESTADOS_TRAMITE_PROVEEDOR)
                        obs = st.text_area("Observaciones del Área de Registro")
                        if st.form_submit_button("Avanzar a Paso 4"):
                            conn.cursor().execute("UPDATE productos SET proveedor=?, tipo_documento=?, numero_documento_oc=?, tramite_proveedor=?, fecha_paso3=?, observacion_paso3=?, paso_actual=4 WHERE id=?", (proveedor, tipo_doc, num_doc, tramite, str(datetime.now().date()), obs, prod_id))
                            conn.commit(); st.success("Avanzado."); time.sleep(1.5); st.rerun()

        with tab_p3_seguimiento:
            df_p3_seg = pd.read_sql_query("SELECT id AS ID, codigo_reyimen AS Código, descripcion AS Descripción, lote AS Lote, vencimiento AS Vencimiento, motivo_informe AS Motivo, proveedor AS Proveedor, tramite_proveedor AS Trámite, numero_documento_oc AS Doc FROM productos WHERE paso_actual >= 3 AND estado_global = 'En trámite' AND motivo_informe != 'Alerta Sanitaria'", conn)
            
            if df_p3_seg.empty: 
                st.info("No hay trámites en seguimiento.")
            else:
                df_p3_seg["Nivel Alerta"] = df_p3_seg.apply(lambda row: calcular_semaforo_vencimiento(row["Vencimiento"], row["Motivo"]), axis=1)
                
                st.markdown("### 🔍 Filtros de Búsqueda")
                col_f1, col_f2 = st.columns(2)
                filtro_semaforo_seg = col_f1.multiselect("Filtrar por Semáforo de Riesgo", ["🔴 Roja (≤ 60d)", "🟡 Amarilla (61-120d)", "🟢 Verde (> 120d)"], placeholder="Seleccione colores...", key="sem_p3_s")
                filtro_codigo_seg = col_f2.text_input("Filtrar por Código Reyimen", placeholder="Ej: 543...", key="cod_p3_s")
                
                df_filtrado_seg = df_p3_seg.copy()
                if filtro_semaforo_seg: df_filtrado_seg = df_filtrado_seg[df_filtrado_seg["Nivel Alerta"].isin(filtro_semaforo_seg)]
                if filtro_codigo_seg: df_filtrado_seg = df_filtrado_seg[df_filtrado_seg["Código"].astype(str).str.contains(filtro_codigo_seg, case=False, na=False)]
                
                st.markdown("<br>", unsafe_allow_html=True)
                if df_filtrado_seg.empty:
                    st.warning("No hay trámites que coincidan con los filtros.")
                else:
                    cols_mostrar_seg = ["ID", "Nivel Alerta", "Código", "Descripción", "Proveedor", "Trámite", "Doc"]
                    st.dataframe(df_filtrado_seg[cols_mostrar_seg], hide_index=True, use_container_width=True, height=180)
                    
                    opciones_id_seg = df_filtrado_seg['ID'].tolist()
                    formato_opciones_seg = {row['ID']: f"{row['Código']} - {row['Descripción']} (Lote: {row['Lote']})" for idx, row in df_filtrado_seg.iterrows()}
                    
                    st.markdown("### ⚙️ Actualización del Trámite")
                    id_seg = st.selectbox("Seleccione el producto filtrado para actualizar:", opciones_id_seg, format_func=lambda x: formato_opciones_seg[x], key="sel_p3_s")
                    
                    prod_seg = conn.cursor().execute("SELECT * FROM productos WHERE id=?", (id_seg,)).fetchone()
                    if prod_seg:
                        with st.form("form_paso3_actualizar"):
                            u_prov = st.selectbox("Proveedor", PROVEEDORES_OFICIALES, index=PROVEEDORES_OFICIALES.index(prod_seg['proveedor']) if prod_seg['proveedor'] in PROVEEDORES_OFICIALES else 0)
                            u_doc = st.text_input("N° Doc", value=prod_seg['numero_documento_oc'] or "")
                            u_tram = st.selectbox("Estado", ESTADOS_TRAMITE_PROVEEDOR, index=ESTADOS_TRAMITE_PROVEEDOR.index(prod_seg['tramite_proveedor']) if prod_seg['tramite_proveedor'] in ESTADOS_TRAMITE_PROVEEDOR else 0)
                            u_obs = st.text_area("Obs.", value=prod_seg['observacion_paso3'] or "")
                            if st.form_submit_button("Guardar"):
                                conn.cursor().execute("UPDATE productos SET proveedor=?, numero_documento_oc=?, tramite_proveedor=?, observacion_paso3=? WHERE id=?", (u_prov, u_doc, u_tram, u_obs, id_seg))
                                conn.commit(); st.success("Actualizado."); time.sleep(1); st.rerun()

    # --- PASO 4 ---
    elif tab_seleccionada == "📦 4. Bulto/Ubicación":
        st.header("📦 Paso 4 — Bulto y Ubicaciones")
        tab_p4_nuevo, tab_p4_seg = st.tabs(["➕ Asignar Nueva Ubicación", "🔄 Seguimiento de Bultos"])
        
        with tab_p4_nuevo:
            df = pd.read_sql_query("SELECT id AS ID, codigo_reyimen AS Código, descripcion AS Descripción, lote AS Lote, proveedor AS Proveedor, vencimiento AS Vencimiento, motivo_informe AS Motivo FROM productos WHERE paso_actual = 4 AND estado_global = 'En trámite' AND motivo_informe != 'Alerta Sanitaria'", conn)
            
            if df.empty: 
                st.info("No hay productos pendientes de asignación de bulto.")
            else:
                df["Nivel Alerta"] = df.apply(lambda row: calcular_semaforo_vencimiento(row["Vencimiento"], row["Motivo"]), axis=1)
                
                st.markdown("### 🔍 Filtros de Búsqueda")
                col_f1, col_f2 = st.columns(2)
                filtro_semaforo = col_f1.multiselect("Filtrar por Semáforo de Riesgo", ["🔴 Roja (≤ 60d)", "🟡 Amarilla (61-120d)", "🟢 Verde (> 120d)"], placeholder="Seleccione colores...", key="sem_p4_n")
                filtro_codigo = col_f2.text_input("Filtrar por Código Reyimen", placeholder="Ej: 543...", key="cod_p4_n")
                
                df_filtrado = df.copy()
                if filtro_semaforo: df_filtrado = df_filtrado[df_filtrado["Nivel Alerta"].isin(filtro_semaforo)]
                if filtro_codigo: df_filtrado = df_filtrado[df_filtrado["Código"].astype(str).str.contains(filtro_codigo, case=False, na=False)]
                
                st.markdown("<br>", unsafe_allow_html=True)
                if df_filtrado.empty:
                    st.warning("No hay productos que coincidan con los filtros.")
                else:
                    cols_mostrar = ["ID", "Nivel Alerta", "Código", "Descripción", "Lote", "Proveedor"]
                    st.dataframe(df_filtrado[cols_mostrar], hide_index=True, use_container_width=True, height=180)
                    
                    opciones_id = df_filtrado['ID'].tolist()
                    formato_opciones = {row['ID']: f"{row['Código']} - {row['Descripción']} (Lote: {row['Lote']})" for idx, row in df_filtrado.iterrows()}
                    
                    st.markdown("### ⚙️ Gestión de Bulto")
                    prod_id = st.selectbox("Seleccione el producto filtrado:", opciones_id, format_func=lambda x: formato_opciones[x], key="sel_p4_n")
                    
                    with st.form("form_p4"):
                        ub_fisica, ub_comp = st.selectbox("Ubicación Física *", OPCIONES_FISICA_P4), st.selectbox("Ubicación Computacional *", BODEGAS_PASO4)
                        bulto, obs = st.text_input("N° Bulto *"), st.text_area("Observaciones Paso 4")
                        if st.form_submit_button("Avanzar a Paso 5"):
                            conn.cursor().execute("UPDATE productos SET ubicacion_fisica=?, ubicacion_computacional=?, numero_bulto=?, observacion_paso4=?, paso_actual=5 WHERE id=?", (ub_fisica, ub_comp, bulto, obs, prod_id))
                            conn.commit(); st.success("Avanzado a Paso 5."); time.sleep(1); st.rerun()
                        
        with tab_p4_seg:
            df_seg4 = pd.read_sql_query("SELECT id AS ID, codigo_reyimen AS Código, descripcion AS Descripción, lote AS Lote, ubicacion_fisica AS Física, numero_bulto AS Bulto, vencimiento AS Vencimiento, motivo_informe AS Motivo FROM productos WHERE paso_actual >= 4 AND estado_global = 'En trámite' AND motivo_informe != 'Alerta Sanitaria'", conn)
            
            if df_seg4.empty: 
                st.info("No hay bultos activos en seguimiento.")
            else:
                df_seg4["Nivel Alerta"] = df_seg4.apply(lambda row: calcular_semaforo_vencimiento(row["Vencimiento"], row["Motivo"]), axis=1)
                
                st.markdown("### 🔍 Filtros de Búsqueda")
                col_f1, col_f2 = st.columns(2)
                filtro_semaforo_seg = col_f1.multiselect("Filtrar por Semáforo de Riesgo", ["🔴 Roja (≤ 60d)", "🟡 Amarilla (61-120d)", "🟢 Verde (> 120d)"], placeholder="Seleccione colores...", key="sem_p4_s")
                filtro_codigo_seg = col_f2.text_input("Filtrar por Código Reyimen", placeholder="Ej: 543...", key="cod_p4_s")
                
                df_filtrado_seg = df_seg4.copy()
                if filtro_semaforo_seg: df_filtrado_seg = df_filtrado_seg[df_filtrado_seg["Nivel Alerta"].isin(filtro_semaforo_seg)]
                if filtro_codigo_seg: df_filtrado_seg = df_filtrado_seg[df_filtrado_seg["Código"].astype(str).str.contains(filtro_codigo_seg, case=False, na=False)]
                
                st.markdown("<br>", unsafe_allow_html=True)
                if df_filtrado_seg.empty:
                    st.warning("No hay bultos que coincidan con los filtros.")
                else:
                    cols_mostrar_seg = ["ID", "Nivel Alerta", "Código", "Descripción", "Física", "Bulto"]
                    st.dataframe(df_filtrado_seg[cols_mostrar_seg], hide_index=True, use_container_width=True, height=180)
                    
                    opciones_id_seg = df_filtrado_seg['ID'].tolist()
                    formato_opciones_seg = {row['ID']: f"{row['Código']} - {row['Descripción']} (Lote: {row['Lote']})" for idx, row in df_filtrado_seg.iterrows()}
                    
                    st.markdown("### ⚙️ Actualización de Ubicación")
                    id_seg4 = st.selectbox("Seleccione el producto filtrado para actualizar:", opciones_id_seg, format_func=lambda x: formato_opciones_seg[x], key="sel_p4_s")
                    
                    prod_seg4 = conn.cursor().execute("SELECT * FROM productos WHERE id=?", (id_seg4,)).fetchone()
                    if prod_seg4:
                        with st.form("form_p4_actualizar"):
                            idx_fis = OPCIONES_FISICA_P4.index(prod_seg4['ubicacion_fisica']) if prod_seg4['ubicacion_fisica'] in OPCIONES_FISICA_P4 else 0
                            idx_comp = BODEGAS_PASO4.index(prod_seg4['ubicacion_computacional']) if prod_seg4['ubicacion_computacional'] in BODEGAS_PASO4 else 0
                            
                            u_fisica = st.selectbox("Ubicación Física", OPCIONES_FISICA_P4, index=idx_fis)
                            u_comp = st.selectbox("Ubicación Computacional", BODEGAS_PASO4, index=idx_comp)
                            u_bulto = st.text_input("N° Bulto", value=prod_seg4['numero_bulto'] or "")
                            u_obs4 = st.text_area("Observaciones", value=prod_seg4['observacion_paso4'] or "")
                            
                            if st.form_submit_button("Actualizar Ubicación/Bulto"):
                                conn.cursor().execute("UPDATE productos SET ubicacion_fisica=?, ubicacion_computacional=?, numero_bulto=?, observacion_paso4=? WHERE id=?", (u_fisica, u_comp, u_bulto, u_obs4, id_seg4))
                                conn.commit(); st.success("Ubicación actualizada."); time.sleep(1); st.rerun()

    # --- PASO 5 ---
    elif tab_seleccionada == "📜 5. Resolución/Cierre":
        st.header("📜 Paso 5 — Resolución y Cierre")
        tab_p5_cierre, tab_p5_sin_canje = st.tabs(["🔒 Cierre General", "🟢 Sin Carta de Canje"])
        
        with tab_p5_cierre:
            df_p5 = pd.read_sql_query("SELECT id AS ID, codigo_reyimen AS Código, descripcion AS Descripción, lote AS Lote, proveedor AS Proveedor, numero_bulto AS Bulto, estado_global AS Estado, vencimiento AS Vencimiento, motivo_informe AS Motivo FROM productos WHERE paso_actual = 5 AND estado_global != 'Concluido'", conn)
            
            if df_p5.empty: 
                st.info("No hay productos pendientes.")
            else:
                df_p5["Nivel Alerta"] = df_p5.apply(lambda row: calcular_semaforo_vencimiento(row["Vencimiento"], row["Motivo"]), axis=1)
                
                st.markdown("### 🔍 Filtros de Búsqueda")
                col_f1, col_f2 = st.columns(2)
                filtro_semaforo = col_f1.multiselect("Filtrar por Semáforo de Riesgo", ["🔴 Roja (≤ 60d)", "🟡 Amarilla (61-120d)", "🟢 Verde (> 120d)"], placeholder="Seleccione colores...", key="sem_p5_c")
                filtro_codigo = col_f2.text_input("Filtrar por Código Reyimen", placeholder="Ej: 543...", key="cod_p5_c")
                
                df_filtrado = df_p5.copy()
                if filtro_semaforo: df_filtrado = df_filtrado[df_filtrado["Nivel Alerta"].isin(filtro_semaforo)]
                if filtro_codigo: df_filtrado = df_filtrado[df_filtrado["Código"].astype(str).str.contains(filtro_codigo, case=False, na=False)]
                
                st.markdown("<br>", unsafe_allow_html=True)
                if df_filtrado.empty:
                    st.warning("No hay productos que coincidan con los filtros.")
                else:
                    cols_mostrar = ["ID", "Nivel Alerta", "Código", "Descripción", "Lote", "Proveedor", "Bulto"]
                    st.dataframe(df_filtrado[cols_mostrar], hide_index=True, use_container_width=True, height=180)
                    
                    opciones_id = df_filtrado['ID'].tolist()
                    formato_opciones = {row['ID']: f"{row['Código']} - {row['Descripción']} (Lote: {row['Lote']})" for idx, row in df_filtrado.iterrows()}
                    
                    st.markdown("### ⚙️ Cierre de Trámite")
                    prod_id = st.selectbox("Seleccione el producto filtrado para CERRAR:", opciones_id, format_func=lambda x: formato_opciones[x], key="sel_p5_c")
                    
                    with st.form("form_p5"):
                        num_res = st.text_input("N° Resolución o Documento *")
                        # --- NUEVO: Opciones detalladas de cierre ---
                        estado_fin = st.selectbox("Estado Final *", [
                            "Canjeado - reposición del producto", 
                            "Canjeado - nota de credito recibida", 
                            "Producto no canjeado", 
                            "Producto dado de baja",
                            "Retirado por Alerta Sanitaria"
                        ])
                        obs = st.text_area("Resolución")
                        if st.form_submit_button("Finalizar y Archivar"):
                            conn.cursor().execute("UPDATE productos SET resolucion_numero=?, estado_final=?, observacion_paso5=?, estado_global='Concluido' WHERE id=?", (num_res, estado_fin, obs, prod_id))
                            conn.commit(); st.success("Archivado."); time.sleep(1.5); st.rerun()
                            
        with tab_p5_sin_canje:
            df_p5_sc = pd.read_sql_query("SELECT id AS ID, codigo_reyimen AS Código, descripcion AS Descripción, lote AS Lote, estado_canje AS Canje, vencimiento AS Vencimiento, motivo_informe AS Motivo FROM productos WHERE (estado_canje = 'No aplica' OR estado_canje LIKE '%compra propia%') AND estado_global = 'En trámite' AND motivo_informe != 'Alerta Sanitaria'", conn)
            
            if not df_p5_sc.empty:
                df_p5_sc["Nivel Alerta"] = df_p5_sc.apply(lambda row: calcular_semaforo_vencimiento(row["Vencimiento"], row["Motivo"]), axis=1)
                
                st.markdown("### 🔍 Filtros de Búsqueda")
                col_f1, col_f2 = st.columns(2)
                filtro_semaforo_sc = col_f1.multiselect("Filtrar por Semáforo de Riesgo", ["🔴 Roja (≤ 60d)", "🟡 Amarilla (61-120d)", "🟢 Verde (> 120d)"], placeholder="Seleccione colores...", key="sem_p5_sc")
                filtro_codigo_sc = col_f2.text_input("Filtrar por Código Reyimen", placeholder="Ej: 543...", key="cod_p5_sc")
                
                df_filtrado_sc = df_p5_sc.copy()
                if filtro_semaforo_sc: df_filtrado_sc = df_filtrado_sc[df_filtrado_sc["Nivel Alerta"].isin(filtro_semaforo_sc)]
                if filtro_codigo_sc: df_filtrado_sc = df_filtrado_sc[df_filtrado_sc["Código"].astype(str).str.contains(filtro_codigo_sc, case=False, na=False)]
                
                st.markdown("<br>", unsafe_allow_html=True)
                if df_filtrado_sc.empty:
                    st.warning("No hay productos sin canje que coincidan con los filtros.")
                else:
                    cols_mostrar_sc = ["ID", "Nivel Alerta", "Código", "Descripción", "Lote", "Canje"]
                    st.dataframe(df_filtrado_sc[cols_mostrar_sc], hide_index=True, use_container_width=True, height=180)
                    
                    opciones_id_sc = df_filtrado_sc['ID'].tolist()
                    formato_opciones_sc = {row['ID']: f"{row['Código']} - {row['Descripción']} (Lote: {row['Lote']})" for idx, row in df_filtrado_sc.iterrows()}
                    
                    st.markdown("### ⚙️ Gestión a la Red")
                    id_sc = st.selectbox("Seleccione el producto para Difusión:", opciones_id_sc, format_func=lambda x: formato_opciones_sc[x], key="sel_p5_sc")
                    
                    prod_sc = conn.cursor().execute("SELECT * FROM productos WHERE id=?", (id_sc,)).fetchone()
                    if prod_sc:
                        with st.form("form_paso5_sin_canje"):
                            c1, c2 = st.columns(2)
                            with c1: difusion_sel = st.selectbox("DIFUSIÓN A LA RED", OPCIONES_DIFUSION_RED)
                            with c2: redistribucion_sel = st.selectbox("REDISTRIBUCIÓN STOCK", OPCIONES_REDISTRIBUCION_STOCK)
                            obs_sc = st.text_area("Observaciones", value=prod_sc['observacion_paso5'] or "")
                            if st.form_submit_button("Guardar Gestión"):
                                conn.cursor().execute("UPDATE productos SET tipo_gestion_canje=?, observacion_paso2=?, observacion_paso5=? WHERE id=?", (difusion_sel, redistribucion_sel, obs_sc, id_sc))
                                conn.commit(); st.success("Guardado."); time.sleep(1); st.rerun()

    # --- ALERTAS, CARGA MASIVA Y ADMIN ---
    elif tab_seleccionada == "🚨 Gestión Anexo II (Jefatura)":
        st.header("🚨 Gestión de Anexo II (Alertas Sanitarias)")
        df_alertas = pd.read_sql_query("SELECT id AS ID, alerta_numero AS 'N° Alerta', codigo_reyimen AS Código, descripcion AS Descripción, lote AS Lote, cantidad AS 'Cant.', proveedor AS 'Proveedor Asignado', estado_global AS Estado FROM productos WHERE motivo_informe = 'Alerta Sanitaria' AND estado_global IN ('CUARENTENA', 'Alerta Notificada al Proveedor')", conn)
        if df_alertas.empty: st.info("No hay Alertas Sanitarias activas pendientes de gestión.")
        else:
            st.dataframe(df_alertas, hide_index=True, use_container_width=True, height=180)
            id_alerta = st.selectbox("Seleccione ID de Alerta a gestionar", df_alertas['ID'].tolist())
            prod_alerta = conn.cursor().execute("SELECT * FROM productos WHERE id=?", (id_alerta,)).fetchone()
            if prod_alerta:
                st.markdown(f"### Redacción de Anexo II para: {prod_alerta['descripcion']}")
                with st.form("form_anexo_ii"):
                    col1, col2 = st.columns(2)
                    principio_activo = col1.text_input("Principio Activo", value=prod_alerta['principio_activo'] or "")
                    titular = col1.text_input("Titular del Producto *", value=prod_alerta['titular_registro'] or "")
                    reg_sanitario = col1.text_input("N° de Registro Sanitario *", value=prod_alerta['registro_sanitario'] or "")
                    proveedor = col2.selectbox("Proveedor *", PROVEEDORES_OFICIALES, index=PROVEEDORES_OFICIALES.index(prod_alerta['proveedor']) if prod_alerta['proveedor'] in PROVEEDORES_OFICIALES else 0)
                    rep_legal = col2.text_input("Representante Legal Hospital", value="Falta Rellenar")
                    dir_tecnico = col2.text_input("Director Técnico / QF Responsable", value=user_info['nombre_completo'])
                    obs_alerta = st.text_area("Otras Observaciones", value=prod_alerta['observacion_paso2'] or "")
                    if st.form_submit_button("💾 Guardar Datos para Anexo II"):
                        conn.cursor().execute("UPDATE productos SET principio_activo=?, titular_registro=?, registro_sanitario=?, proveedor=?, observacion_paso2=? WHERE id=?", (principio_activo, titular, reg_sanitario, proveedor, obs_alerta, id_alerta))
                        conn.commit(); st.success("Datos guardados."); time.sleep(1); st.rerun()
                if prod_alerta['titular_registro'] and prod_alerta['registro_sanitario']:
                    datos_docx = { "descripcion": prod_alerta['descripcion'], "principio_activo": prod_alerta['principio_activo'], "titular": prod_alerta['titular_registro'], "registro_sanitario": prod_alerta['registro_sanitario'], "lote": prod_alerta['lote'], "representante_legal": "Representante Hospital", "director_tecnico": user_info['nombre_completo'], "proveedor": prod_alerta['proveedor'], "cantidad": prod_alerta['cantidad'], "unidad": prod_alerta['unidad'], "observaciones": prod_alerta['observacion_paso2'] }
                    st.download_button("📄 Descargar Documento Anexo II (.docx)", data=generar_anexo_ii_docx(datos_docx), file_name=f"ANEXO_II_{prod_alerta['lote']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    if prod_alerta['estado_global'] != 'Alerta Notificada al Proveedor':
                        if st.button("✅ Marcar como 'Notificado al Proveedor'"):
                            conn.cursor().execute("UPDATE productos SET estado_global='Alerta Notificada al Proveedor', paso_actual=5 WHERE id=?", (id_alerta,))
                            conn.commit(); st.success("Avanzado."); time.sleep(1.5); st.rerun()

    elif tab_seleccionada == "📤 Carga Masiva":
        st.header("📤 Carga Masiva de Productos (Paso 1)")
        df_plantilla = pd.DataFrame([{"BODEGA ORIGEN": "Bodega AZ09 (Fármacos)", "TIPO PRODUCTO": "Fármaco", "CÓDIGO REYIMEN": "1365", "DESCRIPCIÓN": "BUPIVACAINA", "TIPO COMPRA": "CENABAST", "UNIDAD": "FRASCO", "CANTIDAD": 100, "FECHA VENCIMIENTO": "2026-10-31", "LOTE": "L12345"}])
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer: df_plantilla.to_excel(writer, index=False, sheet_name='Plantilla_Carga')
        st.download_button(label="📥 Descargar Plantilla Excel", data=output.getvalue(), file_name="Plantilla_Carga_Masiva.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        uploaded = st.file_uploader("Subir archivo Excel o CSV", type=["xlsx", "xls", "csv"])
        if uploaded and st.button("🚀 Procesar e Ingresar Productos"):
            ok, msg = procesar_carga_masiva(uploaded, user_info['usuario'])
            if ok: st.success(msg); time.sleep(1.5); st.rerun()
            else: st.error(msg)
            
    elif tab_seleccionada == "🔍 Consolidado General":
        st.header("🔍 Consolidado")
        df_cons = pd.read_sql_query("SELECT id AS ID, codigo_reyimen AS Código, descripcion AS Descripción, bodega_origen AS Bodega, cantidad AS Cant, lote AS Lote, vencimiento AS Venc, motivo_informe AS Motivo, estado_global AS Estado, proveedor AS Proveedor FROM productos", conn)
        
        df_cons["Nivel Alerta"] = df_cons.apply(lambda row: calcular_semaforo_vencimiento(row["Venc"], row["Motivo"]), axis=1)
        cols_cons = ["ID", "Nivel Alerta", "Código", "Descripción", "Bodega", "Cant", "Lote", "Venc", "Estado", "Proveedor"]
        df_cons = df_cons[cols_cons]
        st.dataframe(df_cons, hide_index=True)
        
    elif tab_seleccionada == "📊 Dashboard / Análisis":
        st.header("📊 Dashboard")
        df_all = pd.read_sql_query("SELECT * FROM productos", conn)
        if not df_all.empty:
            df_all["Nivel Alerta"] = df_all.apply(lambda row: calcular_semaforo_vencimiento(row["vencimiento"], row["motivo_informe"]), axis=1)
            
            rojas = len(df_all[df_all["Nivel Alerta"] == "🔴 Roja (≤ 60d)"])
            amarillas = len(df_all[df_all["Nivel Alerta"] == "🟡 Amarilla (61-120d)"])
            verdes = len(df_all[df_all["Nivel Alerta"] == "🟢 Verde (> 120d)"])
            
            st.markdown("### 🚦 Semáforo de Riesgo (Vencimientos)")
            sem1, sem2, sem3 = st.columns(3)
            sem1.metric("🔴 Alerta Roja (Crítico)", rojas)
            sem2.metric("🟡 Alerta Amarilla (Tramitación)", amarillas)
            sem3.metric("🟢 Alerta Verde (Monitoreo)", verdes)
            
            st.markdown("---")
            
            m1, m2, m3, m4 = st.columns(4)
            m1.metric("Trámites Totales", len(df_all))
            m2.metric("En Trámite Activo", len(df_all[df_all['estado_global'].isin(['En trámite', 'CUARENTENA', 'Alerta Notificada al Proveedor'])]))
            m3.metric("Concluidos", len(df_all[df_all['estado_global'] == 'Concluido']))
            
            # --- NUEVO: Actualización del contador de Unidades Canjeadas ---
            estados_canje = ['Canjeado - reposición del producto', 'Canjeado - nota de credito recibida']
            canjeadas_df = df_all[df_all['estado_final'].isin(estados_canje)]
            unidades_canjeadas = int(canjeadas_df['cantidad'].sum()) if not canjeadas_df.empty else 0
            m4.metric("Unid. Canjeadas", unidades_canjeadas)
            
    elif tab_seleccionada == "👥 Gestión de Usuarios":
        st.header("👥 Gestión de Usuarios")
        tab_crear, tab_editar = st.tabs(["➕ Crear Usuario", "✏️ Editar / Eliminar"])
        with tab_crear:
            with st.form("form_nuevo_usuario"):
                u_user, u_pass = st.text_input("Usuario"), st.text_input("Contraseña", type="password")
                u_nombre, u_rol = st.text_input("Nombre"), st.selectbox("Rol", ["admin", "jefatura", "registro", "bodega"])
                if st.form_submit_button("Crear"):
                    conn.cursor().execute("INSERT INTO usuarios (usuario, password, rol, nombre_completo) VALUES (?, ?, ?, ?)", (u_user, u_pass, u_rol, u_nombre))
                    conn.commit()
                    st.success("Usuario creado."); time.sleep(1); st.rerun()
            st.dataframe(pd.read_sql_query("SELECT id, usuario, rol, nombre_completo, estado FROM usuarios", conn), hide_index=True, use_container_width=True, height=180)
        with tab_editar:
            df_users_edit = pd.read_sql_query("SELECT * FROM usuarios", conn)
            if not df_users_edit.empty:
                id_mod = st.selectbox("ID a Modificar", df_users_edit['id'].tolist())
                user_data = conn.cursor().execute("SELECT * FROM usuarios WHERE id=?", (id_mod,)).fetchone()
                if user_data:
                    with st.form("form_editar_usuario"):
                        new_u_nombre, new_u_user = st.text_input("Nombre", value=user_data['nombre_completo']), st.text_input("Usuario", value=user_data['usuario'])
                        new_u_rol, new_u_estado = st.selectbox("Rol", ["admin", "jefatura", "registro", "bodega"]), st.selectbox("Estado", ["Activo", "Inactivo"])
                        new_u_pass = st.text_input("Nueva Contraseña (Opcional)", type="password")
                        b1, b2 = st.columns(2)
                        if b1.form_submit_button("Guardar"):
                            if new_u_pass.strip(): conn.cursor().execute("UPDATE usuarios SET usuario=?, password=?, rol=?, nombre_completo=?, estado=? WHERE id=?", (new_u_user, new_u_pass, new_u_rol, new_u_nombre, new_u_estado, id_mod))
                            else: conn.cursor().execute("UPDATE usuarios SET usuario=?, rol=?, nombre_completo=?, estado=? WHERE id=?", (new_u_user, new_u_rol, new_u_nombre, new_u_estado, id_mod))
                            conn.commit(); st.success("Actualizado."); time.sleep(1); st.rerun()
                        if b2.form_submit_button("Eliminar"):
                            conn.cursor().execute("DELETE FROM usuarios WHERE id=?", (id_mod,))
                            conn.commit(); st.warning("Eliminado."); time.sleep(1); st.rerun()

    conn.close()
